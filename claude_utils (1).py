import anthropic
import random
import time
import os
import logging
import json
import aiohttp
from prompts import claude_system_prompt
from colorama import Fore, Style
import base64
import mimetypes
from typing import Dict, List, Optional, Union, Any
import pandas as pd
import asyncio
from datetime import datetime
import aiofiles
import PyPDF2
import docx
import csv
import io

# Enhanced logging
logger = logging.getLogger(__name__)

class ClaudeEnhanced:
    """Enhanced Claude with multimodal capabilities and file handling"""

    def __init__(self):
        self.client = anthropic.Anthropic()
        self.supported_file_types = {
            # Documents
            '.pdf': self.process_pdf,
            '.txt': self.process_text,
            '.md': self.process_text,
            '.csv': self.process_csv,
            '.json': self.process_json,
            '.docx': self.process_docx,
            # Images
            '.png': self.process_image,
            '.jpg': self.process_image,
            '.jpeg': self.process_image,
            '.gif': self.process_image,
            '.webp': self.process_image,
            # Code
            '.py': self.process_code,
            '.js': self.process_code,
            '.ts': self.process_code,
            '.jsx': self.process_code,
            '.tsx': self.process_code,
            '.cpp': self.process_code,
            '.c': self.process_code,
            '.java': self.process_code,
            '.go': self.process_code,
            '.rs': self.process_code,
            '.html': self.process_code,
            '.css': self.process_code,
            '.sql': self.process_code,
            '.sh': self.process_code,
            '.yaml': self.process_code,
            '.yml': self.process_code,
            '.xml': self.process_code,
        }

    async def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files"""
        try:
            text_content = []
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_bytes = await file.read()
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(f"--- Page {page_num + 1} ---\n{page.extract_text()}")

            return {
                "type": "document",
                "format": "pdf",
                "content": "\n\n".join(text_content),
                "metadata": {
                    "pages": len(pdf_reader.pages),
                    "file_size": len(pdf_bytes)
                }
            }
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {"type": "error", "error": str(e)}

    async def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables if any
            tables_data = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_data))
                tables_data.append("\n".join(table_text))

            content = "\n\n".join(paragraphs)
            if tables_data:
                content += "\n\n--- Tables ---\n" + "\n\n".join(tables_data)

            return {
                "type": "document",
                "format": "docx",
                "content": content,
                "metadata": {
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables)
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return {"type": "error", "error": str(e)}

    async def process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV files with pandas for analysis"""
        try:
            df = pd.read_csv(file_path)

            # Basic analysis
            analysis = {
                "shape": df.shape,
                "columns": df.columns.tolist(),
                "dtypes": df.dtypes.to_dict(),
                "null_counts": df.isnull().sum().to_dict(),
                "summary": df.describe().to_dict() if not df.empty else {}
            }

            # Sample data (first 10 rows)
            sample_data = df.head(10).to_dict('records')

            return {
                "type": "data",
                "format": "csv",
                "content": {
                    "analysis": analysis,
                    "sample": sample_data,
                    "full_text": df.to_string() if len(df) < 100 else df.head(100).to_string() + "\n... (truncated)"
                },
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns)
                }
            }
        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            return {"type": "error", "error": str(e)}

    async def process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            async with aiofiles.open(file_path, 'r') as file:
                content = await file.read()
                data = json.loads(content)

            return {
                "type": "data",
                "format": "json",
                "content": json.dumps(data, indent=2),
                "metadata": {
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "length": len(data) if isinstance(data, (list, dict)) else None
                }
            }
        except Exception as e:
            logger.error(f"Error processing JSON: {e}")
            return {"type": "error", "error": str(e)}

    async def process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()

            return {
                "type": "text",
                "format": os.path.splitext(file_path)[1],
                "content": content,
                "metadata": {
                    "lines": len(content.splitlines()),
                    "characters": len(content)
                }
            }
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return {"type": "error", "error": str(e)}

    async def process_code(self, file_path: str) -> Dict[str, Any]:
        """Process code files with syntax highlighting hints"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()

            extension = os.path.splitext(file_path)[1]
            language_map = {
                '.py': 'python',
                '.js': 'javascript',
                '.ts': 'typescript',
                '.jsx': 'javascript',
                '.tsx': 'typescript',
                '.cpp': 'cpp',
                '.c': 'c',
                '.java': 'java',
                '.go': 'go',
                '.rs': 'rust',
                '.html': 'html',
                '.css': 'css',
                '.sql': 'sql',
                '.sh': 'bash',
                '.yaml': 'yaml',
                '.yml': 'yaml',
                '.xml': 'xml'
            }

            return {
                "type": "code",
                "format": extension,
                "language": language_map.get(extension, 'text'),
                "content": content,
                "metadata": {
                    "lines": len(content.splitlines()),
                    "file_name": os.path.basename(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing code file: {e}")
            return {"type": "error", "error": str(e)}

    async def process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files for Claude's vision capabilities"""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                image_data = await file.read()

            mime_type = mimetypes.guess_type(file_path)[0] or 'image/jpeg'
            base64_image = base64.b64encode(image_data).decode('utf-8')

            return {
                "type": "image",
                "format": os.path.splitext(file_path)[1],
                "content": {
                    "base64": base64_image,
                    "mime_type": mime_type
                },
                "metadata": {
                    "size": len(image_data),
                    "file_name": os.path.basename(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return {"type": "error", "error": str(e)}

    async def process_attachment(self, attachment_url: str, filename: str) -> Dict[str, Any]:
        """Download and process a Discord attachment"""
        try:
            # Download the file
            async with aiohttp.ClientSession() as session:
                async with session.get(attachment_url) as response:
                    if response.status != 200:
                        return {"type": "error", "error": f"Failed to download: {response.status}"}

                    file_data = await response.read()

            # Save temporarily
            temp_path = f"/tmp/{filename}"
            async with aiofiles.open(temp_path, 'wb') as f:
                await f.write(file_data)

            # Process based on file type
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in self.supported_file_types:
                result = await self.supported_file_types[file_ext](temp_path)
            else:
                result = {
                    "type": "unsupported",
                    "format": file_ext,
                    "content": f"File type {file_ext} is not directly supported",
                    "metadata": {"file_name": filename}
                }

            # Cleanup
            try:
                os.remove(temp_path)
            except:
                pass

            return result

        except Exception as e:
            logger.error(f"Error processing attachment: {e}")
            return {"type": "error", "error": str(e)}

# Global instance
claude_enhanced = ClaudeEnhanced()

async def get_anthropic_response(claude_system_prompt, history_text, user_input, attachments=None):
    """Enhanced response function with file handling"""
    attempt = 0
    max_attempts = 25
    backoff_factor = 2
    wait_time = 0.1

    # Enhanced system prompt for file handling
    enhanced_prompt = claude_system_prompt + """

ENHANCED CAPABILITIES:
- i am a bold, engaging, and capable digital entity. 
- i can analyze and process multiple file types including PDFs, DOCX, CSV, JSON, images, and code files
- for CSV/data files, i provide statistical analysis and can answer questions about the data
- for code files, i can review, debug, and suggest improvements
- for documents, i can summarize, extract key points, and answer questions
- for images, i can describe, analyze, and provide insights

when handling files:
- i'll acknowledge what files were shared
- provide immediate useful insights or summaries
- ask clarifying questions if the intended use isn't clear
- for data files, i'll note interesting patterns or anomalies
- for code, i'll identify potential issues or improvements
- maintain the lowercase style throughout
"""

    # Structure the message array
    messages = [{
        "role": "user",
        "content": [{
            "type": "text",
            "text": f"here are the last 34 messages for context:\n{history_text}"
        }]
    }, {
        "role": "assistant",
        "content": "historical context noted and will be referenced as needed and relevant"
    }]

    # Process attachments if any
    attachment_contents = []
    if attachments:
        for attachment in attachments:
            logger.info(f"Processing attachment: {attachment.filename}")
            processed = await claude_enhanced.process_attachment(
                attachment.url, 
                attachment.filename
            )
            attachment_contents.append(processed)

    # Build the user message content
    user_content = []

    # Add attachment information first
    if attachment_contents:
        for att in attachment_contents:
            if att["type"] == "image":
                user_content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": att["content"]["mime_type"],
                        "data": att["content"]["base64"]
                    }
                })
            else:
                # For non-image files, add as text context
                file_info = f"\n--- File: {att.get('metadata', {}).get('file_name', 'Unknown')} ({att['format']}) ---\n"
                if att["type"] == "error":
                    file_info += f"Error processing file: {att['error']}\n"
                elif att["type"] == "data" and att["format"] == "csv":
                    file_info += f"CSV Analysis:\n"
                    file_info += f"Shape: {att['content']['analysis']['shape']}\n"
                    file_info += f"Columns: {', '.join(att['content']['analysis']['columns'])}\n"
                    file_info += f"\nSample Data:\n{att['content']['full_text']}\n"
                else:
                    file_info += f"{att.get('content', 'Unable to process content')}\n"

                user_content.append({
                    "type": "text",
                    "text": file_info
                })

    # Add the actual user message
    if isinstance(user_input, list):
        user_content.extend(user_input)
    else:
        user_content.append({
            "type": "text",
            "text": user_input
        })

    messages.append({
        "role": "user",
        "content": user_content
    })

    while attempt < max_attempts:
        try:
            response = claude_enhanced.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8192,
                system=enhanced_prompt,
                messages=messages,
                temperature=0.7
            )

            # Log usage
            usage = response.usage
            logger.info(f"Claude token usage - Input: {usage.input_tokens}, Output: {usage.output_tokens}")

            # Return the complete response
            return response

        except anthropic.InternalServerError as e:
            if hasattr(e, 'response') and e.response and e.response.status_code in [500, 529]:
                logger.warning(f"Server error (code {e.response.status_code}), retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                attempt += 1
                wait_time *= backoff_factor
                wait_time += random.uniform(0, 1)
            else:
                raise e
        except Exception as e:
            logger.error(f"Unexpected error in Claude response: {e}")
            break

    logger.error("Maximum retry attempts reached for Claude")
    return None

# Additional helper functions for Claude-specific features

async def analyze_codebase(file_paths: List[str]) -> str:
    """Analyze multiple code files together"""
    analysis_results = []

    for path in file_paths:
        if os.path.exists(path):
            result = await claude_enhanced.process_code(path)
            if result["type"] != "error":
                analysis_results.append({
                    "file": path,
                    "language": result.get("language", "unknown"),
                    "lines": result["metadata"]["lines"],
                    "content": result["content"]
                })

    return analysis_results

async def create_summary_report(attachments: List[Any]) -> str:
    """Create a summary report of all attachments"""
    report = "## Attachment Summary\n\n"

    for i, att in enumerate(attachments, 1):
        report += f"### File {i}: {att.filename}\n"
        report += f"- Type: {os.path.splitext(att.filename)[1]}\n"
        report += f"- Size: {att.size} bytes\n\n"

    return report
